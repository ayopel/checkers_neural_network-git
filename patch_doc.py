#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Patch: add reflection text + fix Hebrew RTL display throughout.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

PATH = "/sessions/beautiful-eloquent-volta/mnt/checkers_neural_network-git/project_summary.docx"

doc = Document(PATH)

# ── RTL fix: patch every existing paragraph ──────────────────
def ensure_rtl(para):
    pPr = para._p.get_or_add_pPr()
    # bidi
    if pPr.find(qn('w:bidi')) is None:
        b = OxmlElement('w:bidi'); pPr.insert(0, b)
    # right-align (skip centered paras)
    if para.alignment not in (WD_ALIGN_PARAGRAPH.CENTER,):
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc'); pPr.append(jc)
        jc.set(qn('w:val'), 'right')
    # patch every run: add cs font so Hebrew glyphs render correctly
    for run in para.runs:
        rPr = run._r.get_or_add_rPr()
        # rFonts – make sure cs (complex script) font is set
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
        current = rFonts.get(qn('w:ascii'), '')
        if not current or current == 'David':
            rFonts.set(qn('w:ascii'),    'David')
            rFonts.set(qn('w:hAnsi'),    'David')
            rFonts.set(qn('w:cs'),       'David')
        else:
            # code/Latin font – still set cs to avoid fallback issues
            if not rFonts.get(qn('w:cs')):
                rFonts.set(qn('w:cs'), current)

for para in doc.paragraphs:
    ensure_rtl(para)

# Also patch inside tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                ensure_rtl(para)

# ── Also set document-level bidi default ─────────────────────
body = doc.element.body
docDefaults = doc.element.find('.//' + qn('w:docDefaults'))
if docDefaults is not None:
    rPrDefault = docDefaults.find('.//' + qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = OxmlElement('w:rPrDefault'); docDefaults.append(rPrDefault)
    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr'); rPrDefault.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.insert(0, rFonts)
    rFonts.set(qn('w:cs'), 'David')

# ── Find reflection section and replace placeholder lines ─────
REFLECTION = [
    ("בכנות? בהתחלה לא הבנתי כמה הפרויקט הזה יהיה גדול. חשבתי שאני אכתוב קצת קוד, "
     "הרשת תלמד לשחק דמקה, וזהו. לקח לי בערך שבוע עד שהבנתי שאני לא יודע כמעט כלום."),

    ("הדבר הראשון שנפל עלי היה שרשת נוירונים זה לא קסם. ישבתי שעות על הנייר וניסיתי "
     "להבין איך בכלל ה-Feed-Forward עובד – למה כופלים, למה מחברים, מה זה בעצם "
     "\"למידה\" אם אין Backpropagation. הרבה יוטיוב, הרבה ויקיפדיה, והרבה רגעים "
     "שפשוט סגרתי את הלפטופ ויצאתי להסתובב כי הראש לא הכיל יותר."),

    ("הדבר שהכי קשה לי היה לבחור אלגוריתם גנטי במקום Backprop. כולם עושים Backprop, "
     "זה מה שלמדנו, זה מה שיש בכל מדריך. אבל כשהתחלתי לחשוב איך מלמדים את הרשת "
     "לשחק דמקה בלי label מוגדר לכל מהלך – הבנתי שזה פשוט לא עובד ככה. אז החלטתי "
     "ללכת על הגנטי, ולא הייתי בטוח שזו החלטה חכמה עד שראיתי את זה עובד."),

    ("הרגע הכי טוב בכל הפרויקט היה כשהרצתי את האימון בפעם הראשונה והאי ניצח בדור "
     "ה-30 יותר ממה שהפסיד. פשוט ישבתי וצפיתי בלוג גולל בצג ולא עשיתי כלום, כי "
     "זה היה מוזר ויפה לראות שמשהו שכתבתי \"מחליט\" בעצמו."),

    ("אבל היו גם הרבה רגעים פחות יפים. ה-Thread Safety היה כאב ראש אמיתי – קיבלתי "
     "פלטים שלא הגיוניים באימון המקבילי, ולא הבנתי למה. שעתיים חיפשתי באגים "
     "בלוגיקה עד שמישהו הסביר לי שזה לא הלוגיקה, זה שכמה תהליכים כותבים לאותו "
     "מקום בו-זמנית. Interlocked פתר את זה בשורה אחת, וזה הרגיש קצת מטופש בדיעבד."),

    ("אם הייתי מתחיל היום, הייתי קודם כל מבין את כל המתמטיקה לפני שאני כותב שורת "
     "קוד אחת. הלכתי הרבה מהר מדי לקוד ואחר כך נאלצתי לחזור ולשכתב דברים כי לא "
     "הבנתי מה אני בונה. גם הייתי כותב unit tests מהתחלה – בדיקת המהלכים החוקיים "
     "בדמקה לקחה ממני ימים שיכלו להיות שעות."),

    ("מה שאני לוקח מהפרויקט הזה זה בעיקר שלא לפחד מדברים שנראים גדולים מדי. "
     "בהתחלה \"רשת נוירונים\" נשמע כמו משהו שרק חוקרים בגוגל עושים. עכשיו אני יודע "
     "בדיוק מה יש בפנים, ביט-ביט, שורה-שורה. זה לא קסם, זה כפל מטריצות ופונקציית "
     "הפעלה. וזה אולי הדבר הכי חשוב שלמדתי."),
]

# Locate the reflection heading paragraph index
refl_start = None
for i, p in enumerate(doc.paragraphs):
    if 'רפלקציה' in p.text and '5.' in p.text:
        refl_start = i
        break

if refl_start is not None:
    # Find the callout paragraph right after the heading, then the blank-line placeholders
    # Strategy: remove placeholder paragraphs (those with just spaces / empty bordered lines)
    # and insert real reflection paragraphs after the callout

    # find the callout para (first non-empty para after heading)
    callout_idx = refl_start + 1
    while callout_idx < len(doc.paragraphs) and doc.paragraphs[callout_idx].text.strip() == '':
        callout_idx += 1

    # The placeholder lined paragraphs come after the callout
    # We'll insert after the callout paragraph in the XML
    callout_para = doc.paragraphs[callout_idx]

    # Build new paragraphs and insert them after callout_para in the XML
    def insert_rtl_para_after(anchor_para, text, size=11, space_after=8):
        """Create a new paragraph in the doc body and insert after anchor."""
        new_p = OxmlElement('w:p')
        # pPr
        pPr = OxmlElement('w:pPr')
        # bidi
        bidi = OxmlElement('w:bidi'); pPr.append(bidi)
        # jc right
        jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'right'); pPr.append(jc)
        # spacing
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), str(int(size * 20 * space_after / 12)))
        spacing.set(qn('w:before'), '40')
        pPr.append(spacing)
        new_p.append(pPr)
        # run
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'David')
        rFonts.set(qn('w:hAnsi'), 'David')
        rFonts.set(qn('w:cs'),    'David')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(size * 2)); rPr.append(sz)
        szCs = OxmlElement('w:szCs'); szCs.set(qn('w:val'), str(size * 2)); rPr.append(szCs)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
        r.append(t)
        new_p.append(r)
        # insert after anchor
        anchor_para._p.addnext(new_p)
        return new_p

    # Insert in reverse order (each one goes right after callout, pushing previous down)
    last_anchor = callout_para
    inserted = []
    for para_text in REFLECTION:
        new_p = insert_rtl_para_after(last_anchor, para_text, size=11, space_after=10)
        # wrap in a Para object to use as next anchor
        from docx.text.paragraph import Paragraph
        last_anchor = Paragraph(new_p, doc)
        inserted.append(last_anchor)

    # Now remove the old blank-line placeholder paragraphs that follow
    # (they have bottom borders and lots of space_after)
    # Collect them: paragraphs after last inserted that are blank/near-blank
    all_paras = list(doc.paragraphs)
    last_new_idx = all_paras.index(last_anchor) if last_anchor in all_paras else -1
    if last_new_idx == -1:
        # fall back: find by xml element
        all_p_elems = list(doc.element.body)
        last_new_idx = all_p_elems.index(last_anchor._p) if last_anchor._p in all_p_elems else -1

    # remove blank placeholder paras following insertion point
    body_children = list(doc.element.body)
    last_p_idx = body_children.index(last_anchor._p)
    to_remove = []
    for el in body_children[last_p_idx+1:]:
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag == 'p':
            text = ''.join(t.text or '' for t in el.findall('.//' + qn('w:t')))
            if text.strip() == '' or text.strip() == ' ':
                to_remove.append(el)
            else:
                break  # stop at next real content
        else:
            break
    for el in to_remove:
        doc.element.body.remove(el)

doc.save(PATH)
print(f"Patched and saved: {PATH}")
