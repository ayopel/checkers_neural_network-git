#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Checkers AI – Professional project-portfolio document (תיק פרויקט)
Design: deep-blue headings, coloured tables, styled code blocks, info callouts.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "/sessions/beautiful-eloquent-volta/mnt/checkers_neural_network-git/project_summary.docx"

# ══════════════════════════════════════════════════════════
# COLOUR PALETTE
# ══════════════════════════════════════════════════════════
C_DARK_BLUE   = RGBColor(0x1F, 0x38, 0x64)   # headings
C_MID_BLUE    = RGBColor(0x2E, 0x75, 0xB6)   # sub-headings, table header bg alt
C_LIGHT_BLUE  = RGBColor(0xD6, 0xE4, 0xF0)   # table alt rows, callout bg
C_DEEP_RED    = RGBColor(0xC0, 0x00, 0x00)   # title accent
C_WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
C_DARK_GRAY   = RGBColor(0x40, 0x40, 0x40)
C_CODE_BG     = "F2F2F2"   # hex string for shading
C_TBL_HEADER  = "1F3864"   # hex for table header fill
C_TBL_ALT     = "D6E4F0"   # hex for alternating rows
C_CALLOUT_BG  = "FFF9E6"   # info box background (pale yellow)
C_DIVIDER     = "2E75B6"   # section rule colour

# ══════════════════════════════════════════════════════════
# LOW-LEVEL XML HELPERS
# ══════════════════════════════════════════════════════════

def _set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    # remove old shading if present
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    tcPr.append(shd)

def _set_para_shading(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    for old in pPr.findall(qn('w:shd')):
        pPr.remove(old)
    pPr.append(shd)

def _set_para_border(para, side, color_hex, size=12, space=1):
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    el = OxmlElement(f'w:{side}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), str(size))
    el.set(qn('w:space'), str(space))
    el.set(qn('w:color'), color_hex)
    # remove old
    for old in pBdr.findall(qn(f'w:{side}')):
        pBdr.remove(old)
    pBdr.append(el)

def _set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    # bidi
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.insert(0, bidi)
    # right align
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'right')

def _run(para, text, bold=False, italic=False, size=None, color=None, name='David'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = name
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def _page_number_footer(section):
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    run.font.size = Pt(10)
    run.font.name = 'David'
    for tag, text in [('begin', None), (None, ' PAGE '), ('end', None)]:
        if tag == 'begin':
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin')
            run._r.append(fc)
        elif tag == 'end':
            fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'end')
            run._r.append(fc)
        else:
            it = OxmlElement('w:instrText'); it.text = text
            run._r.append(it)

# ══════════════════════════════════════════════════════════
# HIGH-LEVEL PARAGRAPH BUILDERS
# ══════════════════════════════════════════════════════════

def rtl_para(doc, text, bold=False, size=12, italic=False,
             color=None, center=False, spacing_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl(p)
    pf = p.paragraph_format
    pf.space_after  = Pt(spacing_after)
    pf.space_before = Pt(2)
    _run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p

def h1(doc, text):
    """Main section heading – dark blue, 16pt, bold, with bottom border."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl(p)
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after  = Pt(6)
    _set_para_border(p, 'bottom', C_DIVIDER, size=8, space=4)
    run = _run(p, text, bold=True, size=16, color=C_DARK_BLUE)
    return p

def h2(doc, text):
    """Sub-heading – mid-blue, 13pt bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl(p)
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after  = Pt(4)
    _run(p, text, bold=True, size=13, color=C_MID_BLUE)
    return p

def h3(doc, text):
    """Tertiary heading – dark gray, 12pt bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl(p)
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after  = Pt(3)
    _run(p, text, bold=True, size=12, color=C_DARK_GRAY)
    return p

def bullet(doc, text, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl(p)
    pf = p.paragraph_format
    pf.space_after  = Pt(3)
    pf.left_indent  = Cm(0.8)
    # manual bullet
    run = p.add_run('►  ' + text)   # ► bullet
    run.font.size = Pt(size)
    run.font.name = 'David'
    return p

def code_block(doc, code_text):
    """Shaded code block with left border."""
    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_after  = Pt(0)
        pf.space_before = Pt(0)
        pf.left_indent  = Cm(0.5)
        _set_para_shading(p, C_CODE_BG)
        _set_para_border(p, 'left', '2E75B6', size=16, space=6)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x20, 0x20, 0x60)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def callout(doc, title, body):
    """Pale-yellow info box with left accent border."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl(p)
    _set_para_shading(p, C_CALLOUT_BG)
    _set_para_border(p, 'left', 'FFC000', size=24, space=6)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after  = Pt(2)
    pf.left_indent  = Cm(0.4)
    _run(p, f'{title}: ', bold=True, size=11, color=RGBColor(0xC0, 0x60, 0x00))
    _run(p, body, bold=False, size=11, color=C_DARK_GRAY)
    return p

def divider(doc):
    p = doc.add_paragraph()
    _set_para_border(p, 'bottom', C_DIVIDER, size=4, space=1)
    p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════
# TABLE HELPERS
# ══════════════════════════════════════════════════════════

def styled_table(doc, headers, rows, col_widths_cm=None):
    """
    headers: list of str
    rows: list of list of str
    Alternating row shading, dark blue header.
    """
    ncols = len(headers)
    nrows = len(rows) + 1
    tbl = doc.add_table(rows=nrows, cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        _set_cell_shading(cell, C_TBL_HEADER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = C_WHITE
        run.font.size = Pt(10)
        run.font.name = 'David'

    # Data rows
    for i, row_data in enumerate(rows):
        fill = C_TBL_ALT if i % 2 == 0 else 'FFFFFF'
        for j, cell_text in enumerate(row_data):
            cell = tbl.rows[i+1].cells[j]
            _set_cell_shading(cell, fill)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.name = 'David'

    # Column widths
    if col_widths_cm:
        for row in tbl.rows:
            for j, w in enumerate(col_widths_cm):
                row.cells[j].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl

# ══════════════════════════════════════════════════════════
# DOCUMENT SETUP
# ══════════════════════════════════════════════════════════

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

# Default Normal style
style = doc.styles['Normal']
style.font.name = 'David'
style.font.size = Pt(12)

STUDENT = "שם התלמיד"
PROJECT  = "Checkers AI – Deep Neural Network"

# Header + footer for every section
for section in doc.sections:
    hdr = section.header
    hp  = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_para_border(hp, 'bottom', C_DIVIDER, size=6, space=3)
    _run(hp, f'{STUDENT}  │  {PROJECT}', size=9, color=C_DARK_GRAY)
    _page_number_footer(section)

# ══════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════

# Banner background block (simulate with a shaded paragraph)
for _ in range(3):
    p = doc.add_paragraph()
    _set_para_shading(p, '1F3864')
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)

p_title_banner = doc.add_paragraph()
_set_para_shading(p_title_banner, '1F3864')
p_title_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title_banner.paragraph_format.space_before = Pt(0)
p_title_banner.paragraph_format.space_after  = Pt(0)
_run(p_title_banner, 'Checkers AI', bold=True, size=40, color=C_WHITE, name='Arial')

p_sub_banner = doc.add_paragraph()
_set_para_shading(p_sub_banner, '1F3864')
p_sub_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub_banner.paragraph_format.space_before = Pt(0)
p_sub_banner.paragraph_format.space_after  = Pt(0)
_run(p_sub_banner, 'Deep Neural Network Edition', bold=False, size=18,
     color=RGBColor(0xBD, 0xD7, 0xEE), name='Arial')

for _ in range(3):
    p = doc.add_paragraph()
    _set_para_shading(p, '1F3864')
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)

doc.add_paragraph().paragraph_format.space_after = Pt(16)

# Subtitle
rtl_para(doc, 'תיק פרויקט', bold=True, size=22, center=True, color=C_DARK_BLUE)
rtl_para(doc, 'חלופת למידת מכונה – הנדסת תוכנה 883589', size=13, center=True, color=C_DARK_GRAY)

doc.add_paragraph().paragraph_format.space_after = Pt(10)
divider(doc)
doc.add_paragraph().paragraph_format.space_after = Pt(10)

# Info table on cover
cover_data = [
    ['שם התלמיד', '[שם מלא]'],
    ['ת.ז.', '[מספר ת.ז.]'],
    ['שם המנחה', '[שם המנחה]'],
    ['שם החלופה', 'למידת מכונה – הנדסת תוכנה'],
    ['בית הספר', '[שם בית הספר]'],
    ['תאריך הגשה', '[תאריך]'],
]
styled_table(doc, ['פרט', 'ערך'], cover_data, col_widths_cm=[5, 9])

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════
h1(doc, 'תוכן עניינים')
toc_items = [
    ('1.', 'מבוא'),
    ('2.', 'מבנה / ארכיטקטורה'),
    ('', '2.1  שלב איסוף והכנת הנתונים'),
    ('', '2.2  שלב בניית המודל ואימונו'),
    ('', '2.3  שלב היישום'),
    ('3.', 'מדריך למפתח'),
    ('4.', 'מדריך למשתמש'),
    ('5.', 'רפלקציה / סיכום אישי'),
    ('6.', 'ביבליוגרפיה'),
    ('7.', 'נספחים'),
]
for num, name in toc_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl(p)
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    indent = Cm(0.8) if not num else Cm(0)
    pf.left_indent = indent
    if num:
        _run(p, f'{num}  ', bold=True, size=12, color=C_MID_BLUE)
        _run(p, name, bold=True, size=12, color=C_DARK_BLUE)
    else:
        _run(p, name, size=11, color=C_DARK_GRAY, italic=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════
h1(doc, '1. מבוא')

h2(doc, 'רקע הפרויקט')
rtl_para(doc,
    'פרויקט זה מממש משחק דמקה (Checkers) עם בינה מלאכותית מבוססת רשת נוירונים עמוקה '
    '(Deep Neural Network). המערכת מאמנת את הבינה המלאכותית באמצעות אלגוריתם גנטי '
    'וסימולציה עצמית (Self-Play) – ללא צורך בבסיס נתונים חיצוני של מהלכים אנושיים. '
    'הפרויקט ממומש ב-C# עם ממשק גרפי Windows Forms.')
rtl_para(doc,
    'מטרת הפרויקט: לפתח מנגנון למידת מכונה שיוכל לשחק דמקה ברמה גבוהה, תוך הדגמה '
    'מעשית של עקרונות ליבה כגון חישוב קדמה (Feed-Forward), אלגוריתם גנטי, ופונקציית כושר '
    '(Fitness Function) – בלי שימוש בספריות AI חיצוניות.')

h2(doc, 'קהל היעד')
rtl_para(doc,
    'הפרויקט מיועד לכל המעוניין ללמוד על רשתות נוירונים ולמידת חיזוק דרך דוגמה מעשית '
    'ויישום שלם. ניתן להשתמש בו כפלטפורמת לימוד ומחקר בתחום ה-AI.')

h2(doc, 'שלבי הפרויקט')
bullet(doc, 'שלב 1 – בניית מנוע המשחק: לוח, כלים, חוקי מהלכים, תנאי ניצחון ותיקו.')
bullet(doc, 'שלב 2 – בניית ואימון הרשת הנוירונית: ארכיטקטורה, חשמול קדמה, אלגוריתם גנטי.')
bullet(doc, 'שלב 3 – פיתוח ממשק המשתמש: תפריט ראשי, מסך משחק, מסך אימון עם גרף בזמן אמת.')

h2(doc, 'תהליך המחקר')
rtl_para(doc,
    'לפני הפיתוח נסקר המצב הקיים בתחום: מנועי שחמט (Stockfish), AlphaGo/AlphaZero, '
    'וספריות כגון TensorFlow ו-PyTorch. בשל מגבלות סביבת C# ו-Windows Forms הוחלט לממש '
    'את הרשת הנוירונית מאפס (from scratch) ללא תלויות חיצוניות.')
bullet(doc, 'חידוש מרכזי: אימון עצמי מלא ללא נתוני אנוש – האלגוריתם מגלה אסטרטגיות בכוחות עצמו.')
bullet(doc, 'שימוש באלגוריתם גנטי במקום Backpropagation – מתאים לבעיות שקשה לגזור מהן פונקציית שגיאה.')
bullet(doc, 'ריצה מקבילית (Parallel Processing) לזירוז שלב האימון.')
bullet(doc, 'מימוש ידני של He Initialization, Leaky ReLU, Crossover ו-Mutation.')

h2(doc, 'אתגרים מרכזיים')
bullet(doc, 'מניעת קיפאון (Stagnation) – מוטציה אדפטיבית וגיוון גנטי לאחר 20 דורות ללא שיפור.')
bullet(doc, 'בטיחות ריבוי-תהליכים (Thread Safety) – כל עדכוני הסטטיסטיקות מבוצעים עם Interlocked.')
bullet(doc, 'ייצוג לוח יעיל – State String ומטמון מהלכים (ConcurrentDictionary) לביצועים.')
bullet(doc, 'שקלול פונקציית הכושר – איזון עדין בין שיעור ניצחון, ניהול כלים ומלכים.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ══════════════════════════════════════════════════════════
h1(doc, '2. מבנה / ארכיטקטורה של הפרויקט')
rtl_para(doc,
    'הפרויקט מחולק לשלושה שלבים בהתאם לדרישות חלופת למידת מכונה: '
    'איסוף והכנת הנתונים, בניית המודל ואימונו, ופיתוח היישום.')

# ── 2.1 Data ──────────────────────────────────────────────
h2(doc, '2.1  שלב איסוף, הכנה וניתוח הנתונים')
rtl_para(doc,
    'בפרויקט זה אין Dataset חיצוני – הנתונים נוצרים בזמן אמת דרך סימולציה של משחקים. '
    'בכל דור, כל AI משחק נגד מספר יריבים, ותוצאות המשחקים מהוות את ה"נתונים" לאימון.')

h3(doc, 'מבנה הנתון – ייצוג מצב הלוח (Board State)')
rtl_para(doc,
    'הלוח מיוצג כמערך שטוח של 64 ערכים (8×8), הממופה שורה-אחר-שורה. '
    'כל תא מקבל ערך כמותי המקודד גם את זהות הכלי וגם את ערכו האסטרטגי:')
bullet(doc, '0.0 – תא ריק')
bullet(doc, '+ערך חיובי – כלי של השחקן הנוכחי (1.0 עד ~5.0, לפי מיקום וסוג)')
bullet(doc, '-ערך שלילי – כלי של היריב (אותם ערכים עם סימן הפוך)')

callout(doc, 'הערה',
        'כלים שחורים מוצבים בשורות 0-2 (למעלה) וכלים אדומים בשורות 5-7 (למטה). '
        'שורת המלוכה לאדום היא שורה 0, ולשחור שורה 7.')

h3(doc, 'נרמול הנתונים – חישוב ערך כלי')
rtl_para(doc, 'ערך כל כלי מורכב ממספר רכיבים משוקללים המוגדרים בקוד כקבועים:')
styled_table(doc,
    ['רכיב', 'משקל/ערך', 'הסבר'],
    [
        ['ערך בסיסי', 'רגיל=1.0, מלך=3.0', 'בסיס לכל חישוב'],
        ['Center Control', '×0.4', 'גרדיאנט אוקלידי מהמרכז (1.0 במרכז, 0.0 בפינה)'],
        ['Advancement', '×0.3', 'כמה שורות התקדם לכיוון שורת המלוכה'],
        ['Promotion Proximity', '×0.5', 'קרבה להמלכה – ריבועי (אקספוננציאלי)'],
        ['Diagonal Control', '×0.2', 'שליטה באלכסונים הראשיים (+0.3 לכל אלכסון)'],
        ['Edge Penalty', '-0.15', 'עונש לכלים בעמודות 0 ו-7 (פגיעות)'],
        ['Back Row Bonus', '+0.2', 'בונוס לכלים בשורה האחורית (הגנה)'],
        ['Mobility Bonus', '+0.2', 'בונוס אם קיימים מהלכים אפשריים'],
        ['Protection Bonus', '+0.3', 'בונוס אם הכלי מוגן על-ידי כלי אחר'],
    ],
    col_widths_cm=[3.8, 3.0, 8.0]
)

# ── 2.2 Model ─────────────────────────────────────────────
h2(doc, '2.2  שלב בניית המודל ואימון')

h3(doc, 'ארכיטקטורת הרשת הנוירונית')
rtl_para(doc,
    'הרשת היא רשת נוירונים עמוקה מסוג Fully Connected (DNN) '
    'המוגדרת ב-DeepNeuralNetwork.cs:')
styled_table(doc,
    ['שכבה', 'גודל', 'פונקציית הפעלה', 'הערה'],
    [
        ['קלט  (Input)', '64', '–', 'מצב הלוח המקודד'],
        ['שכבה נסתרת 1', '128', 'Leaky ReLU', 'רחבה ביותר – לולד תכונות'],
        ['שכבה נסתרת 2', '64', 'Leaky ReLU', 'דחיסה – הפשטה גבוהה יותר'],
        ['שכבה נסתרת 3', '32', 'Leaky ReLU', 'ייצוג מרוכז'],
        ['פלט  (Output)', '1', 'ליניארי', 'ציון עמדה רציף'],
    ],
    col_widths_cm=[4.0, 2.5, 3.5, 5.5]
)
callout(doc, 'ארכיטקטורה', '64 → 128 → 64 → 32 → 1  |  אלגוריתם: Genetic + Self-Play')

h3(doc, 'אתחול המשקולות – He Initialization')
rtl_para(doc,
    'המשקולות מאותחלות לפי חלוקה נורמלית עם סטיית תקן √(2/n) (שיטת He, 2015), '
    'המתאימה לפונקציות הפעלה מסוג ReLU ומונעת נעלם/פיצוץ גרדיאנטים. '
    'ההטיות (Biases) מאותחלות ל-0.01.')
code_block(doc,
    '// He Initialization – DeepNeuralNetwork.cs\n'
    'double stdDev = Math.Sqrt(2.0 / rows);\n'
    'matrix[i][j] = NextGaussian() * stdDev;   // Gaussian ~ N(0, stdDev)')

h3(doc, 'חישוב קדמה – Feed-Forward')
rtl_para(doc,
    'בכל קריאה ל-FeedForward(), מצב הלוח המקודד עובר דרך כל שכבות הרשת. '
    'לכל שכבה (פרט לפלט) מחושב הפלט ואז מופעל Leaky ReLU:')
code_block(doc,
    '// ApplyLayer – DeepNeuralNetwork.cs\n'
    'output[j] = biases[j];\n'
    'for (int i = 0; i < inputs.Length; i++)\n'
    '    output[j] += inputs[i] * layerWeights[i][j];\n'
    '\n'
    'if (useActivation)   // Leaky ReLU\n'
    '    output[j] = (output[j] > 0) ? output[j] : 0.01 * output[j];')
rtl_para(doc,
    'שכבת הפלט אינה משתמשת בפונקציית הפעלה – הערך הסופי הוא ציון רציף '
    'המייצג את "טיב" העמדה מנקודת מבט השחקן הנוכחי.')

h3(doc, 'פונקציית הכושר – Fitness Function')
rtl_para(doc,
    'במקום Backpropagation, האלגוריתם הגנטי מעריך כל AI דרך פונקציית כושר '
    'המחשבת ניקוד מצטבר לאחר דור של משחקים. '
    'לאחר 20 משחקים, משקל הניצחון גדל אדפטיבית (400→1000):')
styled_table(doc,
    ['רכיב', 'משקל', 'מקסימום נקודות'],
    [
        ['שיעור ניצחון (Win Rate)', '400 – 1000  (אדפטיבי)', '1000'],
        ['פנלטי הפסד (Loss Rate)', '× 0.5 ממשקל הניצחון', '500-'],
        ['פנלטי תיקו (Draw Rate)', '× 0.15 ממשקל הניצחון', '150-'],
        ['מאזן כלים (Material)', '× 5.0 לכל כלי', '~100'],
        ['KingsMade', '× 10.0 לכל מלך', '~120'],
        ['KingsCaptured', '× 20.0 לכל לכידה', '~240'],
        ['KingsLost', '× 25.0 (פנלטי)', '~300-'],
        ['שיעור לכידה (CaptureRate)', '× 200.0', '~200'],
        ['שיעור שרידות (SurvivalRate)', '× 50.0', '50'],
    ],
    col_widths_cm=[5.5, 4.5, 4.5]
)

h3(doc, 'האלגוריתם הגנטי – Genetic Algorithm')
bullet(doc, 'אתחול (Init): יצירת אוכלוסייה ראשונית של 50 רשתות נוירונים אקראיות.')
bullet(doc, 'טורניר (Tournament): כל AI משחק נגד עד 5 יריבים, 2 משחקים לכל זוג (שחלוף צבעים).')
bullet(doc, 'בחירה (TournamentSelect): 5 מתמודדים אקראיים, הטוב ביותר עובר.')
bullet(doc, 'הכלאה (Crossover): לכל משקל וקבוע הטיה – 50% סיכוי מהורה 1 או הורה 2.')
bullet(doc, 'מוטציה (Mutate): Gaussian Noise אדפטיבי, Clamping בין -5.0 ל-5.0 למשקולות.')
bullet(doc, 'אליטיזם (Elitism): 10% הטובים ביותר עוברים ישירות לדור הבא ללא שינוי.')
bullet(doc, 'גיוון (InjectDiversity): אחרי 20 דורות ללא שיפור – 20% חלשים נחלפים בחדשים, 20% מקבלים מוטציה חזקה (0.3).')

h3(doc, 'היפר-פרמטרים')
styled_table(doc,
    ['פרמטר', 'ערך ברירת מחדל', 'טווח אפשרי', 'השפעה'],
    [
        ['Population Size', '50', '10 – 200', 'גיוון גנטי'],
        ['Generations', '100', '1 – 1000', 'עומק אימון'],
        ['Mutation Rate', '0.10', '0.01 – 0.50', 'חקירה vs ניצול'],
        ['Elite %', '10%', '5% – 30%', 'שימור פתרונות טובים'],
        ['Opponents / Player', '5', '3 – 20', 'איכות הערכה'],
        ['Games / Pair', '2', '1 – 5', 'שקלול צבעים'],
        ['Max Moves / Game', '200', '–', 'מניעת לולאות אינסוף'],
        ['Diversity Threshold', '20 דורות', '–', 'איתור קיפאון'],
    ],
    col_widths_cm=[4.0, 3.8, 3.5, 4.5]
)

h3(doc, 'מוטציה אדפטיבית – Adaptive Mutation')
rtl_para(doc,
    'קצב המוטציה גדל אוטומטית כשה-AI חלש יחסית לטוב ביותר בדור, '
    'ובמצב קיפאון (generationsWithoutImprovement > 5):')
code_block(doc,
    '// CalculateAdaptiveMutationRate – TrainingSystem.cs\n'
    'double weakness = 1.0 - (avgParentFitness / bestFitness);\n'
    'rate += weakness * 0.15;\n'
    'if (generationsWithoutImprovement > 5)\n'
    '    rate += 0.05 * (generationsWithoutImprovement / 5.0);\n'
    'return Math.Min(rate, 0.5);   // cap at 50%')

h3(doc, 'התמודדות עם הטיה ושונות (Bias & Variance)')
styled_table(doc,
    ['בעיה', 'גורם', 'פתרון בקוד'],
    [
        ['הטיה (Bias)', 'מוטציה חזקה מדי בשלב מוקדם', 'מוטציה אדפטיבית – חלשה בהתחלה, גדלה עם הזמן'],
        ['שונות (Variance)', 'התכנסות לאסטרטגיה אחת', 'InjectDiversity: 20% חדשים + מוטציה חזקה'],
        ['קיפאון', 'אוכלוסייה הומוגנית', 'DiversityResetThreshold = 20 דורות'],
        ['Over-fitting', 'שחקן מנצח רק נגד פרופיל אחד', 'SelectOpponent: 40% יריב אקראי, 60% דומה'],
    ],
    col_widths_cm=[3.5, 5.0, 7.0]
)

# ── 2.3 Deployment ────────────────────────────────────────
h2(doc, '2.3  שלב היישום (Software Deployment)')
rtl_para(doc,
    'הפרויקט ממומש כיישום Windows Forms הכולל שלושה מסכים עיקריים: '
    'תפריט ראשי (MainMenuForm), מסך משחק (CheckersForm) ומסך אימון (TrainingForm).')

h3(doc, 'כיצד היישום משתמש במודל')
rtl_para(doc,
    'לאחר שהרשת אומנה, היא נשמרת לקובץ בינארי (best_checkers_ai.dat) '
    'ונטענת בעת משחק מול אדם. בחירת המהלך מתבצעת כך:')
bullet(doc, 'סימולציה של לוח עתידי לכל מהלך אפשרי (Board.Clone() + ApplyMove()).')
bullet(doc, 'המרת הלוח הדמוי למערך 64 ערכים (GetBoardState()).')
bullet(doc, 'חישוב קדמה (Brain.FeedForward()) – ציון הרשת.')
bullet(doc, 'ציון אסטרטגי (EvaluateStrategy) × 0.15 + טקטי (EvaluateTactics) × 0.1 + מבנה (EvaluateFormation) × 0.08.')
bullet(doc, 'המהלך עם הציון הגבוה ביותר נבחר.')
code_block(doc,
    '// EvaluateMove – AIPlayer.cs\n'
    'double neuralScore    = Brain.FeedForward(boardState)[0];\n'
    'double strategicScore = EvaluateStrategy(board, simBoard, move, color);\n'
    'double tacticalScore  = EvaluateTactics(simBoard, color);\n'
    'double formationScore = EvaluateFormation(simBoard, color);\n'
    'return neuralScore + strategicScore*0.15 + tacticalScore*0.1 + formationScore*0.08;')

h3(doc, 'תרשים זרימת מסכים (Screen Flow)')
styled_table(doc,
    ['מסך מקור', 'פעולה', 'מסך יעד'],
    [
        ['MainMenuForm', 'Play vs Human', 'CheckersForm  (HumanVsHuman, AI=null)'],
        ['MainMenuForm', 'Play vs AI', 'CheckersForm  (HumanVsAI, AI=loaded)'],
        ['MainMenuForm', 'Train New AI', 'TrainingForm'],
        ['TrainingForm', 'Stop & Save / Complete', 'MainMenuForm'],
        ['CheckersForm', 'Close / Game Over', 'MainMenuForm'],
    ],
    col_widths_cm=[4.5, 4.0, 7.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 3. DEVELOPER GUIDE
# ══════════════════════════════════════════════════════════
h1(doc, '3. מדריך למפתח')
rtl_para(doc, 'מבנה תיקיות הפרויקט:')
styled_table(doc,
    ['תיקייה / קובץ', 'תפקיד'],
    [
        ['AI/AIPlayer.cs',           'שחקן AI – בחירת מהלך, הערכת עמדה, פונקציית כושר'],
        ['AI/DeepNeuralNetwork.cs',  'רשת נוירונים עמוקה מאפס'],
        ['AI/TrainingSystem.cs',     'מנהל מחזור האימון הגנטי'],
        ['Core/GameCore.cs',         'הגדרות ליבה: Board, Piece, Move, MoveValidator'],
        ['Core/GameEngine.cs',       'מנוע המשחק – תורות, Undo, זיהוי ניצחון/תיקו'],
        ['Core/PlayerStats.cs',      'סטטיסטיקות Thread-Safe לכל שחקן AI'],
        ['UI/CheckersForm.cs',       'מסך המשחק הגרפי'],
        ['UI/MainMenuForm.cs',       'תפריט ראשי'],
        ['UI/TrainingForm.cs',       'מסך האימון עם גרף ולוג בזמן אמת'],
        ['UI/Program.cs',            'נקודת כניסה – Application.Run(new MainMenuForm())'],
    ],
    col_widths_cm=[5.5, 10.0]
)

FILES = [
    {
        'name': 'GameCore.cs', 'path': 'Core/GameCore.cs',
        'role': 'הגדרות ליבה: Enums, Board, Piece, Move ו-MoveValidator.',
        'vars': [
            ('PieceColor', 'צבע הכלי: Red / Black.'),
            ('PieceType',  'סוג הכלי: Regular / King.'),
            ('GameState',  'מצב המשחק: RedTurn, BlackTurn, RedWins, BlackWins.'),
            ('GameConstants', 'BoardSize=8, KingRowRed=0, KingRowBlack=7, InitialPiecesPerPlayer=12.'),
            ('Board.squares', 'Piece[8,8] – מטריצת הלוח.'),
            ('Board.cachedStateString', 'מחרוזת 64 תווים – מאוחסנת במטמון; מסומנת "dirty" בכל שינוי.'),
            ('MoveValidator.moveCache', 'ConcurrentDictionary – מטמון מהלכים בטוח-תהליכים.'),
        ],
        'funcs': [
            ('Board.InitializeBoard()', 'מציבה 12 כלים שחורים (שורות 0-2) ו-12 אדומים (שורות 5-7) על תאים אי-זוגיים.'),
            ('Board.ApplyMove(Move)', 'מזיזה כלי, מסירה קפוצים, ומבצעת המלכה אוטומטית אם הכלי הגיע לשורת המלוכה.'),
            ('Board.GetStateString()', 'מחרוזת 64 תווים (r/R/b/B/.) לזיהוי חזרות ומניעת תיקו.'),
            ('Board.Clone()', 'מחזירה עותק עצמאי מלא – נדרש לסימולציות ל-Undo ולהערכת AI.'),
            ('MoveValidator.GetValidMoves(Piece)', 'מחזירה מהלכים חוקיים – כופה קפיצות אם קיימות (חוק דמקה).'),
            ('MoveValidator.HasAvailableJumps(PieceColor)', 'בודקת אם קיימות קפיצות לצבע נתון – לכפיית חוק הקפיצה.'),
        ],
    },
    {
        'name': 'GameEngine.cs', 'path': 'Core/GameEngine.cs',
        'role': 'מנוע המשחק – ניהול תורות, ביצוע מהלכים, זיהוי ניצחון/תיקו, Undo.',
        'vars': [
            ('stateHistory',       'מילון מצבי לוח → מספר חזרות (תיקו ב-3 פעמים).'),
            ('moveHistory',        'Stack<GameSnapshot> לתמיכה ב-Undo.'),
            ('movesWithoutCapture','מונה מהלכים ללא לכידה – תיקו לאחר 50.'),
            ('mustContinueJumping','דגל שמחייב המשך רצף קפיצות באותו תור.'),
        ],
        'funcs': [
            ('ExecuteMove(Move)',       'מבצע מהלך, שומר Snapshot, מעדכן מונים, ובודק קפיצות נוספות.'),
            ('CheckWinCondition()',     'בודקת האם לשחקן הנוכחי אין כלים או מהלכים, וקובעת ניצחון.'),
            ('UndoMove()',              'שחזור מצב לוח קודם מ-Stack – תומך בשני מצבי המשחק.'),
            ('IsDraw()',                'תיקו בחזרה (≥3 פעמים) או 50 מהלכים ללא לכידה.'),
            ('GetGameStats()',          'מחזיר GameStats עם ספירת כלים, מלכים ומהלכים לתצוגה.'),
        ],
    },
    {
        'name': 'PlayerStats.cs', 'path': 'Core/PlayerStats.cs',
        'role': 'סטטיסטיקות Thread-Safe לכל שחקן AI – מבוסס Interlocked.',
        'vars': [
            ('_gamesPlayed, _wins, _losses, _draws', 'שדות גיבוי int עם Interlocked.'),
            ('_piecesCaptured / _piecesLost',        'מעקב כלים שנלכדו ואבדו.'),
            ('_kingsMade / _kingsCaptured / _kingsLost', 'מעקב מלכים – משוקלל גבוה בפונקציית הכושר.'),
        ],
        'funcs': [
            ('IncrementX() / AddX(int)', 'Interlocked.Increment/Add – בטוח בריצה מקבילית.'),
            ('WinRate { get }',          'Wins / GamesPlayed – מחושב בזמן הקריאה.'),
            ('Reset()',                  'מאפסת כל השדות בין דורות; נקראת ללא מקביליות.'),
        ],
    },
    {
        'name': 'DeepNeuralNetwork.cs', 'path': 'AI/DeepNeuralNetwork.cs',
        'role': 'רשת נוירונים עמוקה מאפס: חשמול קדמה, שיבוט, מוטציה, הכלאה, שמירה/טעינה.',
        'vars': [
            ('weights',     'List<double[][]> – מטריצות משקולות לכל שכבה.'),
            ('biases',      'List<double[]> – וקטורי הטיה לכל שכבה.'),
            ('hiddenSizes', 'int[] = {128, 64, 32} – גדלי שכבות נסתרות.'),
            ('Fitness',     'double – ניקוד הכושר שחושב ע"י AIPlayer.CalculateFitness().'),
            ('random',      'Random – מופע משותף לכל פעולות האקראיות של הרשת.'),
        ],
        'funcs': [
            ('InitializeWeightMatrix(r,c)', 'He Init: stdDev = √(2/rows), ערכים מ-Gaussian.'),
            ('FeedForward(double[] inputs)','חישוב קדמה שכבה-אחר-שכבה; מחזיר double[] (גודל 1).'),
            ('Mutate(rate, strength=0.2)',  'Gaussian Noise אדפטיבי; Clamping: weights בין ±5.0, biases בין ±2.0.'),
            ('Crossover(partner, rate=0.5)','הכלאה מאחידה – לכל משקל 50% מהורה 1 או הורה 2.'),
            ('SaveToFile / LoadFromFile',   'שמירה/טעינה בינארית (BinaryWriter/Reader) של כל המשקולות, ההטיות, Fitness ו-GamesPlayed.'),
        ],
    },
    {
        'name': 'AIPlayer.cs', 'path': 'AI/AIPlayer.cs',
        'role': 'שחקן AI – עוטף רשת נוירונים, בוחר מהלכים, מעריך עמדות ומחשב כושר.',
        'vars': [
            ('Brain',            'DeepNeuralNetwork של השחקן.'),
            ('Stats',            'PlayerStats – סטטיסטיקות Thread-Safe.'),
            ('_cachedValidator', 'MoveValidator מאוחסן – נוצר מחדש רק אם הלוח השתנה.'),
            ('CenterWeight=0.4, AdvancementWeight=0.3, PromotionWeight=0.5,\nDiagonalWeight=0.2, EdgePenalty=0.15, BackRowBonus=0.2,\nMobilityBonus=0.2, ProtectionBonus=0.3', 'קבועי שקלול מיקום.'),
        ],
        'funcs': [
            ('ChooseMove(board, moves, color)', 'מעריך כל מהלך ומחזיר את בעל הציון הגבוה ביותר.'),
            ('EvaluateMove(board, move, color)', 'neural + strategy×0.15 + tactics×0.1 + formation×0.08.'),
            ('EvaluateFormation(board, color)',  'קשרים אלכסוניים, שורה אחורית, כלים מתקדמים עם תמיכה, צפיפות.'),
            ('EvaluateEndgame(...)',             'סיום: מרכוז, קרבה ליריב (אגרסיבי) / ניידות (הגנתי), פעילות מלכים.'),
            ('CalculateFitness()',              'ניקוד כושר מצטבר אדפטיבי לאחר דור.'),
            ('Clone / Mutate / Crossover',      'מעטפות ל-Brain.Clone/Mutate/Crossover.'),
        ],
    },
    {
        'name': 'TrainingSystem.cs', 'path': 'AI/TrainingSystem.cs',
        'role': 'מנהל מחזור האימון הגנטי: טורניר, כושר, אבולוציה וגיוון.',
        'vars': [
            ('Population',                    'List<AIPlayer> – כל שחקני הדור הנוכחי.'),
            ('Generation',                    'int – מספר הדור הנוכחי.'),
            ('BestPlayer',                    'AIPlayer בעל הכושר הגבוה ביותר.'),
            ('historicalBestFitness',         'double – כושר שיא לכל הזמנים.'),
            ('generationsWithoutImprovement', 'int – מונה לזיהוי קיפאון; איפוס בכל שיפור.'),
            ('ThreadLocalRandom',             'ThreadLocal<Random> – Random בטוח לכל תהליך.'),
        ],
        'funcs': [
            ('RunGeneration()',       'מחזור דור שלם: Reset → Tournament → Fitness → Evolve → StagnationCheck.'),
            ('RunTournament()',       'Parallel.ForEach לריצת כל הזוגות בו-זמנית.'),
            ('SelectOpponent(...)',   '60% יריב בדרגה דומה (±3), 40% אקראי לגיוון.'),
            ('EvolvePopulation()',    'אליטיזם (10%) + TournamentSelect + Breed (Crossover + AdaptiveMutate).'),
            ('InjectDiversity()',     'מחליפה 20% חלשים ומוטציה 0.3 ל-20% נוספים.'),
            ('GetGenerationReport()', 'מחרוזת דוח לדור: Fitness/WinRate/ממוצע + אזהרת קיפאון.'),
        ],
    },
    {
        'name': 'CheckersForm.cs', 'path': 'UI/CheckersForm.cs',
        'role': 'מסך המשחק הגרפי – לוח כפתורים, ניהול תורות, AI אסינכרוני, Undo.',
        'vars': [
            ('boardButtons[8,8]', 'מטריצת Button – כל כפתור מייצג תא בלוח.'),
            ('game',              'GameEngine – מנוע המשחק הפעיל.'),
            ('aiPlayer',          'AIPlayer – null במצב Human vs Human.'),
            ('isAIThinking',      'bool – חוסם קלט משתמש בזמן שה-AI מחשב.'),
            ('lastMoveFrom/To',   'Position? – המהלך האחרון לצביעה כתומה בלוח.'),
        ],
        'funcs': [
            ('Square_Click(...)',     'בחירת כלי / ביצוע מהלך; מפעיל MakeAIMove() בתור ה-AI.'),
            ('MakeAIMove()',          'async – מריץ AI עם השהיות ויזואליות; תומך ברצף קפיצות מרובות.'),
            ('UpdateBoard()',         'מרענן צבעי כפתורים, הדגשות, מהלכים חוקיים וסטטיסטיקות.'),
            ('UndoButton_Click(...)', 'מחזיר 2 מהלכים (שחקן + AI) ב-HumanVsAI, 1 ב-HumanVsHuman.'),
        ],
    },
    {
        'name': 'MainMenuForm.cs', 'path': 'UI/MainMenuForm.cs',
        'role': 'תפריט ראשי – ניתוב למשחק, AI ואימון; בודק קיום קובץ AI.',
        'vars': [('AI_FILE', '"best_checkers_ai.dat" – שם קובץ ה-AI השמור.')],
        'funcs': [
            ('PlayHuman()', 'פותח CheckersForm(HumanVsHuman, null).'),
            ('PlayAI()',    'טוען AI מקובץ (DeepNeuralNetwork.LoadFromFile) ופותח CheckersForm.'),
            ('TrainAI()',   'פותח TrainingForm.'),
        ],
    },
    {
        'name': 'TrainingForm.cs', 'path': 'UI/TrainingForm.cs',
        'role': 'מסך האימון – הגדרות, שליטה (Start/Pause/Stop), גרף ולוג צבעוני בזמן אמת.',
        'vars': [
            ('trainingSystem',             'TrainingSystem – האובייקט הפעיל.'),
            ('fitnessHistory/avgHistory',  'List<double> – היסטוריית כושר לציור הגרף.'),
            ('isTraining / isPaused',      'bool – דגלי מצב לשליטה בריצה.'),
            ('stagnationCount',            'int – מונה דורות ללא שיפור לצורך לוג.'),
        ],
        'funcs': [
            ('BtnStart_Click(...)',     'בונה TrainingConfig ומפעיל RunTraining ב-Task.Run().'),
            ('RunTraining(config, n)',  'לולאת אימון: RunGeneration לכל דור + Checkpoint כל 20 דורות.'),
            ('ChartPanel_Paint(...)',   'ציור גרף: כחול = כושר טוב ביותר, אדום = ממוצע אוכלוסייה.'),
            ('AppendLog(msg, color)',   'הוספת שורה ללוג עם prefix לפי צבע (⚠/✓/💾/❌).'),
            ('SaveBestPlayer()',        'שומר Brain.SaveToFile("best_checkers_ai.dat").'),
        ],
    },
]

for f in FILES:
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    h2(doc, f'קובץ: {f["name"]}')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl(p)
    _run(p, 'מיקום: ', bold=True, size=11, color=C_MID_BLUE)
    _run(p, f['path'], size=11, italic=True, color=C_DARK_GRAY, name='Courier New')

    rtl_para(doc, f'תפקיד: {f["role"]}', size=11)

    h3(doc, 'משתנים עיקריים')
    for vn, vd in f['vars']:
        bullet(doc, f'{vn}  –  {vd}')

    h3(doc, 'פונקציות עיקריות')
    for fn, fd in f['funcs']:
        bullet(doc, f'{fn}  –  {fd}')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 4. USER GUIDE
# ══════════════════════════════════════════════════════════
h1(doc, '4. מדריך למשתמש')

h2(doc, 'דרישות התקנה')
styled_table(doc,
    ['דרישה', 'פרטים'],
    [
        ['מערכת הפעלה', 'Windows 10 / 11'],
        ['.NET Runtime', '.NET Framework 4.7.2 ומעלה (או .NET 6+)'],
        ['ספריות חיצוניות', 'אין – הפרויקט עצמאי לחלוטין'],
        ['קובץ AI', 'best_checkers_ai.dat (נוצר לאחר אימון)'],
    ],
    col_widths_cm=[5, 10.5]
)

h2(doc, 'הפעלה')
rtl_para(doc, 'הרץ את checkers_neural_network.exe. יפתח מסך התפריט הראשי.')

h2(doc, 'מסך 1 – תפריט ראשי (MainMenuForm)')
rtl_para(doc, 'המסך הראשי מכיל שלושה כפתורים:')
styled_table(doc,
    ['כפתור', 'פעולה', 'תנאי'],
    [
        ['"Play vs Human"', 'משחק בין שני שחקנים אנושיים על אותו מחשב', 'תמיד זמין'],
        ['"Play vs AI"',    'משחק מול הבינה המלאכותית (שחקן = אדום, AI = שחור)', 'דורש best_checkers_ai.dat'],
        ['"Train New AI"',  'פתיחת מסך האימון ליצירת AI חדש', 'תמיד זמין'],
    ],
    col_widths_cm=[4, 7, 4.5]
)
callout(doc, 'הודעת מצב', '✓ Trained AI Available – קובץ קיים ומוכן.  ⚠ No trained AI yet – יש לאמן תחילה.')

h2(doc, 'מסך 2 – משחק (CheckersForm)')
rtl_para(doc, 'הלוח מוצג כרשת 8×8 של כפתורים:')
styled_table(doc,
    ['אלמנט', 'משמעות'],
    [
        ['●', 'כלי רגיל'],
        ['♔', 'מלך'],
        ['הדגשה זהב', 'כלי נבחר כרגע'],
        ['הדגשה ירוקה', 'מהלך חוקי אפשרי'],
        ['הדגשה כתומה', 'המהלך האחרון שבוצע'],
        ['🔄 New Game', 'איפוס המשחק (אישור נדרש)'],
        ['↶ Undo Move', 'חזרה: 1 מהלך (H vs H) / 2 מהלכים (H vs AI)'],
        ['פאנל ימין', 'כלים פעילים, מלכים ולכידות – מתעדכן בזמן אמת'],
    ],
    col_widths_cm=[4.5, 11.0]
)
callout(doc, 'חלון קופץ', 'בסיום המשחק: מוצגת הודעת ניצחון ושאלה האם לשחק שוב.')

h2(doc, 'מסך 3 – אימון (TrainingForm)')
rtl_para(doc, 'הגדרות אימון (ניתן לשנות לפני התחלה):')
styled_table(doc,
    ['הגדרה', 'ערך מומלץ', 'הסבר'],
    [
        ['Population Size', '50', 'גודל האוכלוסייה הגנטית'],
        ['Generations', '100', 'מספר הדורות'],
        ['Mutation Rate', '0.10', 'קצב מוטציה (0–0.5)'],
        ['Elite %', '0.10', 'אחוז שורדים ישירים'],
        ['Opponents / Player', '5', 'יריבים לכל שחקן בדור'],
        ['Parallel Processing', 'מופעל', 'מהיר יותר; מנצל ריבוי ליבות'],
    ],
    col_widths_cm=[4.5, 3.5, 7.5]
)
rtl_para(doc, 'כפתורי שליטה: Start Training → Pause / Resume → Stop & Save.')
rtl_para(doc, 'גרף: קו כחול = כושר הטוב ביותר, קו אדום = ממוצע אוכלוסייה.')
callout(doc, 'Checkpoint', 'נשמר אוטומטית כל 20 דורות לקובץ checkpoint_genXX.dat.')

h2(doc, 'מגבלות ידועות')
bullet(doc, 'אין המשך אימון מנקודת עצירה – רק Checkpoints ידניים.')
bullet(doc, 'ה-AI בוחר מהלך בעומק 1 (ללא Alpha-Beta Pruning).')
bullet(doc, 'הלוח ממומש ב-Windows Forms – לא מגיב במהירות מלאה בזמן האימון.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 5. REFLECTION
# ══════════════════════════════════════════════════════════
h1(doc, '5. רפלקציה / סיכום אישי')
callout(doc, 'הנחיה',
        'יש לכתוב לפחות חצי עמוד. כלול: איך הייתה העבודה, מה קיבלת, אילו כלים אתה לוקח, '
        'מה הקשיים, מה המסקנות, ומה היית עושה אחרת.')
doc.add_paragraph().paragraph_format.space_after = Pt(8)
for _ in range(18):      # placeholder lines
    p = doc.add_paragraph()
    _set_para_border(p, 'bottom', 'CCCCCC', size=2, space=1)
    p.paragraph_format.space_after = Pt(16)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 6. BIBLIOGRAPHY
# ══════════════════════════════════════════════════════════
h1(doc, '6. ביבליוגרפיה')
refs = [
    ('Mitchell, M.', '1996', 'An Introduction to Genetic Algorithms', 'MIT Press', ''),
    ('Goodfellow, I., Bengio, Y., & Courville, A.', '2016', 'Deep Learning', 'MIT Press', 'https://www.deeplearningbook.org'),
    ('He, K., Zhang, X., Ren, S., & Sun, J.', '2015', 'Delving Deep into Rectifiers: Surpassing Human-Level Performance on ImageNet Classification', 'arXiv', 'arXiv:1502.01852'),
    ('Microsoft', '2024', 'System.Threading.Interlocked Class', 'Microsoft Docs', 'https://docs.microsoft.com/en-us/dotnet/api/system.threading.interlocked'),
    ('Wikipedia', '2024', 'Draughts (Checkers)', 'Wikipedia', 'https://en.wikipedia.org/wiki/Draughts'),
    ('Silver, D. et al.', '2017', 'Mastering Chess and Shogi by Self-Play with a General Reinforcement Learning Algorithm', 'arXiv', 'arXiv:1712.01815'),
]
for i, (author, year, title, source, url) in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl(p)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.first_line_indent = Cm(-1.2)
    _run(p, f'{i}. ', bold=True, size=11, color=C_MID_BLUE)
    _run(p, f'{author} ({year}). ', bold=False, size=11)
    _run(p, f'{title}. ', italic=True, size=11)
    _run(p, f'{source}.', size=11)
    if url:
        _run(p, f'  {url}', size=10, color=C_MID_BLUE)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 7. APPENDICES
# ══════════════════════════════════════════════════════════
h1(doc, '7. נספחים')

h2(doc, "נספח א' – Leaky ReLU")
rtl_para(doc,
    'פונקציית הפעלה הדומה ל-ReLU אך מאפשרת שיפוע קטן (0.01) לערכים שליליים. '
    'זה מונע את תופעת "מוות הנוירונים" (Dying ReLU) שבה נוירון מפסיק להתעדכן לחלוטין:')
code_block(doc,
    '// DeepNeuralNetwork.cs\n'
    'private double LeakyReLU(double x) => x > 0 ? x : 0.01 * x;\n'
    '\n'
    '// x > 0  → f(x) = x          (רגיל)\n'
    '// x ≤ 0  → f(x) = 0.01 * x   (שיפוע קטן – לא מת)')

h2(doc, "נספח ב' – He Initialization")
rtl_para(doc,
    'שיטת He (2015) מאתחלת משקולות לפי חלוקה נורמלית עם סטיית תקן √(2/n), '
    'שם n הוא מספר הנוירונים בשכבה הקודמת. מניעת Vanishing/Exploding Gradients:')
code_block(doc,
    '// DeepNeuralNetwork.cs – InitializeWeightMatrix()\n'
    'double stdDev = Math.Sqrt(2.0 / rows);   // He: σ = √(2/n_in)\n'
    'matrix[i][j] = NextGaussian() * stdDev;\n'
    '\n'
    '// NextGaussian() – Box-Muller transform:\n'
    'double u1 = 1.0 - random.NextDouble();\n'
    'double u2 = 1.0 - random.NextDouble();\n'
    'return Math.Sqrt(-2.0 * Math.Log(u1)) * Math.Sin(2.0 * Math.PI * u2);')

h2(doc, "נספח ג' – Board State Encoding")
rtl_para(doc, 'קידוד מצב הלוח למחרוזת 64 תווים לזיהוי חזרות ומניעת תיקו:')
code_block(doc,
    '// Board.cs – GetStateString()\n'
    "// '.' = empty\n"
    "// 'r' = Red regular,  'R' = Red King\n"
    "// 'b' = Black regular, 'B' = Black King\n"
    'chars[idx++] = piece == null ? \'.\' :\n'
    '    piece.Color == PieceColor.Red\n'
    '        ? (piece.Type == PieceType.King ? \'R\' : \'r\')\n'
    '        : (piece.Type == PieceType.King ? \'B\' : \'b\');')

h2(doc, "נספח ד' – Fitness Function (קטע קוד מלא)")
code_block(doc,
    '// AIPlayer.cs – CalculateFitness()\n'
    'double experienceMultiplier = Math.Min(1.0, gamesPlayed / 20.0); // cap at 20\n'
    'double winRateWeight = 400 + (600 * experienceMultiplier);       // 400 → 1000\n'
    '\n'
    'fitness += winRate  * winRateWeight;              // Win:  0 – 1000\n'
    'fitness -= lossRate * (winRateWeight * 0.5);      // Loss: 0 – 500\n'
    'fitness -= drawRate * (winRateWeight * 0.15);     // Draw: 0 – 150\n'
    '\n'
    'fitness += materialBalance * 5.0;                 // כלים: ±5 לכלי\n'
    'fitness += Stats.KingsMade     * 10.0;\n'
    'fitness += Stats.KingsCaptured * 20.0;\n'
    'fitness -= Stats.KingsLost     * 25.0;\n'
    '\n'
    'if (Stats.TotalMoves > 0)\n'
    '    fitness += (Stats.PiecesCaptured / (double)Stats.TotalMoves) * 200.0;\n'
    '\n'
    'double survivalRate = 1.0 - (Stats.PiecesLost / (double)(gamesPlayed * 12));\n'
    'fitness += survivalRate * 50.0;\n'
    '\n'
    'Brain.Fitness = Math.Max(0, fitness);')

h2(doc, "נספח ה' – Adaptive Mutation Rate")
code_block(doc,
    '// TrainingSystem.cs – CalculateAdaptiveMutationRate()\n'
    'double rate = mutationRate;   // ברירת מחדל: 0.10\n'
    '\n'
    '// הגדלה עבור הורים חלשים (חקירה > ניצול)\n'
    'double avgParentFitness = (parent1.Brain.Fitness + parent2.Brain.Fitness) / 2.0;\n'
    'double weakness = 1.0 - (avgParentFitness / bestFitness);\n'
    'rate += weakness * 0.15;\n'
    '\n'
    '// הגדלה נוספת בקיפאון\n'
    'if (generationsWithoutImprovement > 5)\n'
    '    rate += 0.05 * (generationsWithoutImprovement / 5.0);\n'
    '\n'
    'return Math.Min(rate, 0.5);   // מקסימום 50%')

h2(doc, "נספח ו' – InjectDiversity")
rtl_para(doc, 'פונקציית הגיוון מופעלת לאחר 20 דורות ללא שיפור:')
code_block(doc,
    '// TrainingSystem.cs – InjectDiversity()\n'
    'var sortedPop = Population.OrderBy(p => p.Brain.Fitness).ToList();\n'
    'int replaceCount = (int)(populationSize * 0.2);  // 20% חלשים\n'
    '\n'
    '// החלפה בשחקנים אקראיים חדשים\n'
    'for (int i = 0; i < replaceCount; i++)\n'
    '    sortedPop[i] = new AIPlayer(random);\n'
    '\n'
    '// מוטציה חזקה (0.3) ל-20% נוספים\n'
    'for (int i = replaceCount; i < replaceCount*2 && i < sortedPop.Count; i++)\n'
    '    sortedPop[i].Mutate(0.3);')

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
