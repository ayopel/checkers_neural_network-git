#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "/sessions/beautiful-eloquent-volta/mnt/checkers_neural_network-git/project_summary.docx"

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.insert(0, bidi)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    return p

def add_rtl_paragraph(doc, text, bold=False, size=12, italic=False, color=None, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'David'
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(p)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'David'
    return p

def add_code(doc, code):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def add_header_footer(doc, student_name, project_name):
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(f'{student_name}  |  {project_name}')
        run.font.size = Pt(10)
        run.font.name = 'David'
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        run.font.size = Pt(10)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = ' PAGE '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'David'
style.font.size = Pt(12)

STUDENT = "שם התלמיד"
PROJECT = "Checkers AI - Deep Neural Network"
add_header_footer(doc, STUDENT, PROJECT)

# COVER
add_rtl_paragraph(doc, "בית הספר - [שם בית הספר]", bold=True, size=14, center=True)
add_rtl_paragraph(doc, " ", size=6, center=True)
add_rtl_paragraph(doc, "תיק פרויקט", bold=True, size=28, center=True)
add_rtl_paragraph(doc, "חלופת למידת מכונה - הנדסת תוכנה 883589", bold=True, size=16, center=True)
add_rtl_paragraph(doc, " ", size=12, center=True)
add_rtl_paragraph(doc, "שם הפרויקט:", bold=True, size=14, center=True)
add_rtl_paragraph(doc, "Checkers AI - Deep Neural Network", size=18, center=True, color=RGBColor(0xB4,0x00,0x00))
add_rtl_paragraph(doc, " ", size=12, center=True)
for label, val in [("שם התלמיד:","[שם מלא]"),("ת.ז.:","[מספר ת.ז.]"),("שם המנחה:","[שם המנחה]"),("שם החלופה:","למידת מכונה - הנדסת תוכנה"),("תאריך הגשה:","[תאריך]")]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{val}  ")
    r1.font.size = Pt(13); r1.font.name = 'David'
    r2 = p.add_run(label)
    r2.bold = True; r2.font.size = Pt(13); r2.font.name = 'David'
doc.add_page_break()

# TOC
add_heading(doc, "תוכן עניינים", level=1)
for item in ["1. מבוא","2. מבנה / ארכיטקטורה","   2.1 שלב איסוף והכנת הנתונים","   2.2 שלב בניית המודל ואימונו","   2.3 שלב היישום","3. מדריך למפתח","4. מדריך למשתמש","5. רפלקציה / סיכום אישי","6. ביבליוגרפיה","7. נספחים"]:
    add_rtl_paragraph(doc, item)
doc.add_page_break()

# 1. INTRO
add_heading(doc, "1. מבוא", level=1)
add_heading(doc, "רקע הפרויקט", level=2)
add_rtl_paragraph(doc, "פרויקט זה מממש משחק דמקה (Checkers) עם בינה מלאכותית מבוססת רשת נוירונים עמוקה (Deep Neural Network). המערכת מאמנת את הבינה המלאכותית באמצעות אלגוריתם גנטי וסימולציה עצמית (Self-Play) ללא צורך בבסיס נתונים חיצוני. הפרויקט כתוב ב-C# עם ממשק גרפי Windows Forms.")
add_rtl_paragraph(doc, "מטרת הפרויקט: לפתח מנגנון למידת מכונה שיוכל לשחק דמקה ברמה גבוהה, תוך הדגמת עקרונות ליבה: חישוב קדמה (Feed-Forward), אלגוריתם גנטי, ופונקציית כושר (Fitness Function).")
add_heading(doc, "קהל היעד", level=2)
add_rtl_paragraph(doc, "הפרויקט מיועד לכל המעוניין ללמוד על רשתות נוירונים ולמידת חיזוק דרך דוגמה מעשית ויישום שלם.")
add_heading(doc, "שלבי הפרויקט", level=2)
add_bullet(doc, "שלב 1 - בניית מנוע המשחק: לוח, כלים, חוקי מהלכים, תנאי ניצחון.")
add_bullet(doc, "שלב 2 - בניית ואימון הרשת הנוירונית: ארכיטקטורה, חשמול קדמה, אלגוריתם גנטי.")
add_bullet(doc, "שלב 3 - פיתוח ממשק המשתמש: תפריט ראשי, מסך משחק, מסך אימון עם גרף.")
add_heading(doc, "תהליך המחקר", level=2)
add_rtl_paragraph(doc, "לפני הפיתוח נסקר המצב הקיים: מנועי שחמט (Stockfish), AlphaGo/AlphaZero, ספריות TensorFlow/PyTorch. בשל מגבלות סביבת C# הוחלט לממש את הרשת מאפס (from scratch).")
add_bullet(doc, "חידוש מרכזי: אימון עצמי מלא ללא נתוני אנוש.")
add_bullet(doc, "שימוש באלגוריתם גנטי במקום Backpropagation.")
add_bullet(doc, "מקביליות (Parallel Processing) לזירוז האימון.")
add_heading(doc, "אתגרים מרכזיים", level=2)
add_bullet(doc, "מניעת קיפאון (Stagnation) - מוטציה אדפטיבית וגיוון גנטי.")
add_bullet(doc, "בטיחות ריבוי-תהליכים (Thread Safety) - Interlocked לכל הסטטיסטיקות.")
add_bullet(doc, "ייצוג לוח יעיל - State String ו-Move Cache לביצועים.")
add_bullet(doc, "שקלול פונקציית הכושר - איזון ניצחון, כלים ומלכים.")
doc.add_page_break()

# 2. ARCHITECTURE
add_heading(doc, "2. מבנה / ארכיטקטורה של הפרויקט", level=1)
add_rtl_paragraph(doc, "הפרויקט מחולק לשלושה שלבים: איסוף הנתונים, בניית המודל, ופיתוח היישום.")

add_heading(doc, "2.1 שלב איסוף, הכנה וניתוח הנתונים", level=2)
add_rtl_paragraph(doc, "בפרויקט זה אין Dataset חיצוני - הנתונים נוצרים בזמן אמת דרך סימולציה של משחקים. בכל דור, כל AI משחק נגד יריבים ותוצאות המשחקים מהוות את ה'נתונים'.")
add_heading(doc, "מבנה הנתון - ייצוג מצב הלוח", level=3)
add_rtl_paragraph(doc, "הלוח מיוצג כמערך של 64 ערכים (8x8). כל תא מקבל ערך כמותי:")
add_bullet(doc, "0.0 - תא ריק")
add_bullet(doc, "+ערך חיובי - כלי של השחקן הנוכחי (1.0 עד ~5.0 לפי מיקום וסוג)")
add_bullet(doc, "-ערך שלילי - כלי של היריב")
add_heading(doc, "נרמול הנתונים - חישוב ערך כלי", level=3)
add_bullet(doc, "ערך בסיסי: כלי רגיל = 1.0, מלך = 3.0")
add_bullet(doc, "שליטה מרכזית (Center Control): גרדיאנט אוקלידי x 0.4")
add_bullet(doc, "קידום (Advancement): מרחק מהמלכה x 0.3")
add_bullet(doc, "קרבה להמלכה (Promotion Proximity): ריבועי x 0.5")
add_bullet(doc, "עונש קצוות (Edge Penalty): -0.15 לעמודות קצה")
add_bullet(doc, "בונוס שורה אחורית (Back Row Bonus): +0.2")
add_bullet(doc, "בונוס ניידות (Mobility): +0.2 אם יש מהלכים")
add_bullet(doc, "בונוס הגנה (Protection): +0.3 אם מוגן")

add_heading(doc, "2.2 שלב בניית המודל ואימון", level=2)
add_heading(doc, "ארכיטקטורת הרשת הנוירונית", level=3)
add_rtl_paragraph(doc, "הרשת היא DNN מסוג Fully Connected עם הארכיטקטורה הבאה:")

tbl = doc.add_table(rows=6, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["שכבה", "גודל", "פונקציית הפעלה"]):
    c = tbl.rows[0].cells[i]
    c.paragraphs[0].add_run(h).bold = True
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for i, row_d in enumerate([
    ["קלט (Input)", "64 נוירונים", "-"],
    ["שכבה נסתרת 1", "128 נוירונים", "Leaky ReLU"],
    ["שכבה נסתרת 2", "64 נוירונים", "Leaky ReLU"],
    ["שכבה נסתרת 3", "32 נוירונים", "Leaky ReLU"],
    ["פלט (Output)", "1 נוירון", "לינארי"],
]):
    for j, txt in enumerate(row_d):
        c = tbl.rows[i+1].cells[j]
        c.paragraphs[0].add_run(txt)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading(doc, "אתחול המשקולות - He Initialization", level=3)
add_rtl_paragraph(doc, "המשקולות מאותחלות לפי חלוקה נורמלית עם סטיית תקן sqrt(2/n) (שיטת He, 2015), המתאימה ל-ReLU. ההטיות מאותחלות ל-0.01.")
add_heading(doc, "חישוב קדמה - Feed-Forward", level=3)
add_rtl_paragraph(doc, "מצב הלוח עובר דרך כל שכבות הרשת. לכל שכבה:")
add_rtl_paragraph(doc, "output[j] = Sum(input[i] x weight[i][j]) + bias[j]", italic=True)
add_rtl_paragraph(doc, "לאחר מכן Leaky ReLU: f(x) = x אם x > 0, ו-0.01x אחרת. שכבת הפלט אינה משתמשת בפונקציית הפעלה - הציון הסופי הוא ייצוג רציף של איכות העמדה.")
add_heading(doc, "פונקציית הכושר - Fitness Function", level=3)
add_rtl_paragraph(doc, "במקום Backpropagation, האלגוריתם הגנטי משתמש בפונקציית כושר:")
add_bullet(doc, "שיעור ניצחון x 400-1000 (אדפטיבי לפי ניסיון)")
add_bullet(doc, "פנלטי הפסד x 0.5 ממשקל הניצחון")
add_bullet(doc, "מאזן כלים (PiecesCaptured - PiecesLost) x 5")
add_bullet(doc, "KingsMade x 10, KingsCaptured x 20, KingsLost x -25")
add_bullet(doc, "שיעור לכידה (CaptureRate) x 200")
add_bullet(doc, "שיעור שרידות (SurvivalRate) x 50")
add_heading(doc, "האלגוריתם הגנטי", level=3)
add_bullet(doc, "אתחול: 50 רשתות נוירונים אקראיות.")
add_bullet(doc, "טורניר: כל AI משחק נגד 5 יריבים, 2 משחקים לזוג.")
add_bullet(doc, "בחירה (TournamentSelect): 5 מתמודדים, הטוב ביותר עובר.")
add_bullet(doc, "הכלאה (Crossover): 50% מכל הורה לכל משקל.")
add_bullet(doc, "מוטציה (Mutate): Gaussian Noise אדפטיבי, Clamping בין -5.0 ל-5.0.")
add_bullet(doc, "אליטיזם: 10% הטובים ביותר עוברים ישירות.")
add_bullet(doc, "גיוון: אחרי 20 דורות ללא שיפור - 20% נחלפים, 20% מקבלים מוטציה חזקה.")
add_heading(doc, "היפר-פרמטרים", level=3)

hp_tbl = doc.add_table(rows=8, cols=3)
hp_tbl.style = 'Table Grid'
hp_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["פרמטר", "ערך ברירת מחדל", "טווח אפשרי"]):
    c = hp_tbl.rows[0].cells[i]
    c.paragraphs[0].add_run(h).bold = True
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for i, row_d in enumerate([
    ["גודל אוכלוסייה", "50", "10-200"],
    ["מספר דורות", "100", "1-1000"],
    ["קצב מוטציה", "0.10", "0.01-0.50"],
    ["אחוז אליטה", "10%", "5%-30%"],
    ["יריבים לשחקן", "5", "3-20"],
    ["מהלכים מקסימום", "200", "-"],
    ["סף קיפאון", "20 דורות", "-"],
]):
    for j, txt in enumerate(row_d):
        c = hp_tbl.rows[i+1].cells[j]
        c.paragraphs[0].add_run(txt)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

add_heading(doc, "התמודדות עם הטיה ושונות", level=3)
add_bullet(doc, "הטיה (Bias): מוטציה חלשה בשלבים מוקדמים - שמירה על פתרונות טובים.")
add_bullet(doc, "שונות (Variance): מוטציה אדפטיבית גדלה בקיפאון; גיוון גנטי להרחבת חקירה.")
add_bullet(doc, "ממוצע כושר מדווח לצד הטוב ביותר - לזיהוי מוקדם של בעיות.")

add_heading(doc, "2.3 שלב היישום", level=2)
add_rtl_paragraph(doc, "הפרויקט ממומש כיישום Windows Forms עם שלושה מסכים: תפריט ראשי, מסך משחק, ומסך אימון.")
add_heading(doc, "כיצד היישום משתמש במודל", level=3)
add_rtl_paragraph(doc, "לאחר אימון הרשת, היא נשמרת לקובץ בינארי (best_checkers_ai.dat). בעת משחק מול אדם:")
add_bullet(doc, "לכל מהלך אפשרי: סימולציה של לוח עתידי (Board.Clone + ApplyMove).")
add_bullet(doc, "המרת הלוח הדמוי למערך 64 ערכים (GetBoardState).")
add_bullet(doc, "חישוב קדמה (FeedForward) - ציון הרשת.")
add_bullet(doc, "ניקוד אסטרטגי x 0.15 + טקטי x 0.1 + מבנה x 0.08 מצורפים לציון.")
add_bullet(doc, "המהלך עם הציון הגבוה ביותר נבחר.")
add_heading(doc, "UML - זרימת מסכים", level=3)
add_rtl_paragraph(doc, "MainMenuForm --> CheckersForm (HumanVsHuman / HumanVsAI)", italic=True)
add_rtl_paragraph(doc, "MainMenuForm --> TrainingForm --> TrainingSystem --> AIPlayer[]", italic=True)
doc.add_page_break()

# 3. DEVELOPER GUIDE
add_heading(doc, "3. מדריך למפתח", level=1)
add_rtl_paragraph(doc, "מבנה תיקיות הפרויקט:", bold=True)
add_bullet(doc, "AI/  -  AIPlayer.cs, DeepNeuralNetwork.cs, TrainingSystem.cs")
add_bullet(doc, "Core/  -  GameCore.cs, GameEngine.cs, PlayerStats.cs")
add_bullet(doc, "UI/  -  CheckersForm.cs, MainMenuForm.cs, TrainingForm.cs, Program.cs")

files = [
    {
        "name": "GameCore.cs", "path": "Core/GameCore.cs",
        "role": "הגדרות ליבה: Enums, מבנה הלוח (Board), כלים (Piece), מהלכים (Move) ובודק מהלכים (MoveValidator).",
        "vars": [
            ("PieceColor", "צבע הכלי: Red / Black."),
            ("PieceType", "סוג הכלי: Regular / King."),
            ("GameState", "מצב המשחק: RedTurn, BlackTurn, RedWins, BlackWins."),
            ("Board.squares", "מטריצת Piece[8,8] המחזיקה את כל הכלים על הלוח."),
            ("Board.cachedStateString", "מחרוזת 64 תווים של מצב הלוח - מאוחסנת במטמון לביצועים."),
            ("MoveValidator.moveCache", "ConcurrentDictionary של מהלכים - מונע חישובים חוזרים בריבוי-תהליכים."),
        ],
        "funcs": [
            ("Board.InitializeBoard()", "מציבה 12 כלים שחורים (שורות 0-2) ו-12 אדומים (שורות 5-7) על תאים אי-זוגיים."),
            ("Board.ApplyMove(Move)", "מזיזה כלי, מסירה קפוצים, ומבצעת המלכה אוטומטית."),
            ("Board.GetStateString()", "מחרוזת ייחודית של מצב הלוח לזיהוי חזרות ומניעת תיקו."),
            ("Board.Clone()", "מחזירה עותק עצמאי מלא של הלוח לסימולציות."),
            ("MoveValidator.GetValidMoves(Piece)", "מחזירה מהלכים חוקיים תוך כפיית קפיצות חובה."),
            ("MoveValidator.HasAvailableJumps(PieceColor)", "בודקת אם קיימות קפיצות לצבע נתון."),
        ],
    },
    {
        "name": "GameEngine.cs", "path": "Core/GameEngine.cs",
        "role": "מנוע המשחק - מנהל תורות, מבצע מהלכים, מזהה ניצחון/תיקו, תומך ב-Undo.",
        "vars": [
            ("stateHistory", "מילון מצבי לוח לספירת חזרות (תיקו אחרי 3 פעמים)."),
            ("moveHistory", "מחסנית Snapshots לתמיכה ב-Undo."),
            ("movesWithoutCapture", "מונה מהלכים ללא לכידה - תיקו אחרי 50."),
            ("mustContinueJumping", "דגל המחייב המשך רצף קפיצות באותו תור."),
        ],
        "funcs": [
            ("ExecuteMove(Move)", "מבצע מהלך, שומר Snapshot, מעדכן מונים, ובודק קפיצות נוספות."),
            ("CheckWinCondition()", "בודקת האם השחקן הנוכחי נעדר כלים או מהלכים."),
            ("UndoMove()", "שחזור לוח מהמחסנית - תומך ב-Human vs Human ו-Human vs AI."),
            ("IsDraw()", "תיקו בחזרה 3 פעמים, או 50 מהלכים ללא לכידה."),
        ],
    },
    {
        "name": "PlayerStats.cs", "path": "Core/PlayerStats.cs",
        "role": "סטטיסטיקות Thread-Safe לכל שחקן AI.",
        "vars": [
            ("_gamesPlayed, _wins, _losses, _draws", "שדות גיבוי המתעדכנים עם Interlocked."),
            ("_piecesCaptured / _piecesLost", "מעקב כלים שנלכדו ואבדו."),
            ("_kingsMade / _kingsCaptured / _kingsLost", "מעקב מלכים - משוקלל גבוה בפונקציית הכושר."),
        ],
        "funcs": [
            ("IncrementX() / AddX()", "כל המתודות משתמשות ב-Interlocked.Increment/Add."),
            ("WinRate { get }", "מחשבת Wins / GamesPlayed בצורה בטוחה."),
            ("Reset()", "מאפסת את כל השדות בין דורות - נקראת ללא מקביליות."),
        ],
    },
    {
        "name": "DeepNeuralNetwork.cs", "path": "AI/DeepNeuralNetwork.cs",
        "role": "רשת נוירונים עמוקה מאפס: חשמול קדמה, שיבוט, מוטציה, הכלאה, שמירה וטעינה.",
        "vars": [
            ("weights", "List<double[][]> - מטריצות משקולות לכל שכבה."),
            ("biases", "List<double[]> - וקטורי הטיה לכל שכבה."),
            ("hiddenSizes", "מערך גדלי שכבות נסתרות: [128, 64, 32]."),
            ("Fitness", "ניקוד הכושר שחושב ע\"י AIPlayer.CalculateFitness()."),
        ],
        "funcs": [
            ("InitializeWeightMatrix(rows, cols)", "He Initialization - סטיית תקן sqrt(2/rows)."),
            ("FeedForward(double[] inputs)", "חישוב קדמה שכבה אחר שכבה עם Leaky ReLU."),
            ("Mutate(mutationRate, strength)", "Gaussian Noise אדפטיבי, Clamping בין -5.0 ל-5.0."),
            ("Crossover(partner)", "הכלאה מאחידה - 50% מכל הורה."),
            ("SaveToFile / LoadFromFile", "שמירה וטעינה בינארית של משקולות, הטיות וכושר."),
        ],
    },
    {
        "name": "AIPlayer.cs", "path": "AI/AIPlayer.cs",
        "role": "שחקן AI - עוטף רשת נוירונים, בחירת מהלך, הערכת עמדה ופונקציית כושר.",
        "vars": [
            ("Brain", "הרשת הנוירונית של השחקן (DeepNeuralNetwork)."),
            ("Stats", "אובייקט PlayerStats עם סטטיסטיקות המשחק."),
            ("_cachedValidator", "MoveValidator מאוחסן למניעת הקצאות חוזרות."),
            ("CenterWeight, AdvancementWeight, ...", "קבועי שקלול עמדה: 0.4, 0.3, 0.5, 0.2, 0.15, 0.2, 0.3."),
        ],
        "funcs": [
            ("ChooseMove(board, validMoves, color)", "מעריך כל מהלך ומחזיר את הטוב ביותר."),
            ("EvaluateMove(board, move, color)", "רשת x 1.0 + אסטרטגיה x 0.15 + טקטיקה x 0.1 + מבנה x 0.08."),
            ("EvaluateFormation(board, color)", "קשרים אלכסוניים, שורה אחורית, כלים מתקדמים עם תמיכה, צפיפות."),
            ("EvaluateEndgame(...)", "הערכת סיום: מרכוז, קרבה ליריב, פעילות מלכים."),
            ("CalculateFitness()", "ניקוד כושר מצטבר לאחר דור."),
        ],
    },
    {
        "name": "TrainingSystem.cs", "path": "AI/TrainingSystem.cs",
        "role": "מנהל מחזור האימון: טורניר, כושר, אבולוציה וגיוון.",
        "vars": [
            ("Population", "רשימת כל שחקני ה-AI בדור הנוכחי."),
            ("Generation", "מספר הדור הנוכחי."),
            ("BestPlayer", "השחקן בעל הכושר הגבוה ביותר."),
            ("historicalBestFitness", "הכושר הגבוה ביותר שהושג - לזיהוי שיפור."),
            ("generationsWithoutImprovement", "מונה לזיהוי קיפאון."),
            ("ThreadLocalRandom", "Random נפרד לכל תהליך - ללא תחרות."),
        ],
        "funcs": [
            ("RunGeneration()", "מחזור דור שלם: איפוס, טורניר, כושר, אבולוציה, קיפאון."),
            ("RunTournament()", "Parallel.ForEach לריצת כל המשחקים בו-זמנית."),
            ("SelectOpponent(...)", "60% יריב ברמה דומה, 40% אקראי."),
            ("EvolvePopulation()", "אליטיזם + הכלאה + מוטציה אדפטיבית."),
            ("InjectDiversity()", "20% נחלפים, 20% מקבלים מוטציה חזקה."),
        ],
    },
    {
        "name": "CheckersForm.cs", "path": "UI/CheckersForm.cs",
        "role": "מסך המשחק - לוח גרפי, ניהול תורות, מהלכי AI אסינכרוניים, Undo.",
        "vars": [
            ("boardButtons[8,8]", "מטריצת Button המייצגת את הלוח."),
            ("game", "אובייקט GameEngine - מנוע המשחק."),
            ("aiPlayer", "שחקן ה-AI (null במצב Human vs Human)."),
            ("isAIThinking", "דגל שמונע פעולות בזמן שה-AI מחשב."),
            ("lastMoveFrom / lastMoveTo", "מיקומי המהלך האחרון להדגשה."),
        ],
        "funcs": [
            ("Square_Click(...)", "ניהול בחירת כלי ומהלך; מעביר תור ל-AI אם נדרש."),
            ("MakeAIMove()", "async - מריץ ה-AI תוך עדכון ויזואלי; תומך ברצף קפיצות."),
            ("UpdateBoard()", "מרענן את הלוח לפי מצב הלוח, הדגשות ומהלכים חוקיים."),
            ("UndoButton_Click(...)", "מחזיר שני מהלכים (שחקן + AI) במצב HumanVsAI."),
        ],
    },
    {
        "name": "MainMenuForm.cs", "path": "UI/MainMenuForm.cs",
        "role": "תפריט ראשי - ניתוב בין משחק, AI ואימון.",
        "vars": [("AI_FILE", "\"best_checkers_ai.dat\" - שם קובץ ה-AI השמור.")],
        "funcs": [
            ("PlayHuman()", "פותח CheckersForm במצב HumanVsHuman."),
            ("PlayAI()", "טוען קובץ AI ופותח CheckersForm במצב HumanVsAI."),
            ("TrainAI()", "פותח את TrainingForm."),
        ],
    },
    {
        "name": "TrainingForm.cs", "path": "UI/TrainingForm.cs",
        "role": "מסך האימון - הגדרות, שליטה, גרף ולוג בזמן אמת.",
        "vars": [
            ("trainingSystem", "אובייקט TrainingSystem הפעיל."),
            ("fitnessHistory / avgFitnessHistory", "היסטוריית כושר לגרף."),
            ("isTraining / isPaused", "דגלי מצב לשליטה."),
            ("stagnationCount", "מונה דורות ללא שיפור."),
        ],
        "funcs": [
            ("BtnStart_Click(...)", "בונה TrainingConfig ומפעיל RunTraining ב-Task.Run."),
            ("RunTraining(config, generations)", "לולאת אימון: RunGeneration לכל דור + עדכון גרף ולוג."),
            ("ChartPanel_Paint(...)", "גרף כושר: כחול=טוב, אדום=ממוצע."),
            ("SaveBestPlayer()", "שומר ה-AI הטוב לקובץ best_checkers_ai.dat."),
        ],
    },
]

for f in files:
    doc.add_paragraph()
    add_heading(doc, f"קובץ: {f['name']}", level=2)
    add_rtl_paragraph(doc, f"מיקום: {f['path']}", italic=True)
    add_rtl_paragraph(doc, f"תפקיד: {f['role']}")
    add_rtl_paragraph(doc, "משתנים עיקריים:", bold=True)
    for vn, vd in f["vars"]:
        add_bullet(doc, f"{vn} - {vd}")
    add_rtl_paragraph(doc, "פונקציות עיקריות:", bold=True)
    for fn, fd in f["funcs"]:
        add_bullet(doc, f"{fn} - {fd}")

doc.add_page_break()

# 4. USER GUIDE
add_heading(doc, "4. מדריך למשתמש", level=1)
add_heading(doc, "דרישות התקנה", level=2)
add_bullet(doc, "מערכת הפעלה: Windows 10 / 11")
add_bullet(doc, ".NET Framework 4.7.2 ומעלה (או .NET 6+)")
add_bullet(doc, "אין תלות בספריות חיצוניות נוספות")
add_heading(doc, "הפעלה", level=2)
add_rtl_paragraph(doc, "הרץ את checkers_neural_network.exe. יפתח תפריט ראשי.")
add_heading(doc, "מסך 1 - תפריט ראשי (MainMenuForm)", level=2)
add_bullet(doc, '"Play vs Human" - משחק בין שני אנשים על אותו מחשב.')
add_bullet(doc, '"Play vs AI" - משחק מול ה-AI (נדרש best_checkers_ai.dat).')
add_bullet(doc, '"Train New AI" - מסך האימון ליצירת AI חדש.')
add_rtl_paragraph(doc, "הודעת מצב: V Trained AI Available / ! No trained AI yet - יש לאמן תחילה.")
add_heading(doc, "מסך 2 - משחק (CheckersForm)", level=2)
add_bullet(doc, "● = כלי רגיל | מ = מלך")
add_bullet(doc, "הדגשה זהב = כלי נבחר | ירוק = מהלך חוקי | כתום = המהלך האחרון")
add_bullet(doc, "New Game - מאפס משחק (אישור נדרש) | Undo Move - מחזיר מהלך אחורה")
add_bullet(doc, "פאנל ימין: סטטיסטיקות כלים, מלכים ולכידות מתעדכנות בזמן אמת.")
add_rtl_paragraph(doc, "אחרי משחק: חלון קופץ שואל אם לשחק שוב.")
add_heading(doc, "מסך 3 - אימון (TrainingForm)", level=2)
add_bullet(doc, "Population Size: גודל אוכלוסייה (מומלץ 50)")
add_bullet(doc, "Generations: מספר דורות (מומלץ 100)")
add_bullet(doc, "Mutation Rate: קצב מוטציה (מומלץ 0.1)")
add_bullet(doc, "Parallel Processing: מהיר יותר; מומלץ למעבד ריבוי-ליבות.")
add_rtl_paragraph(doc, "Start / Pause / Stop & Save. גרף מתעדכן בזמן אמת. Checkpoint נשמר כל 20 דורות.")
add_heading(doc, "מגבלות ידועות", level=2)
add_bullet(doc, "אין המשך אימון מנקודת עצירה - רק Checkpoints ידניים.")
add_bullet(doc, "ה-AI בוחר מהלך בעומק 1 (ללא Alpha-Beta Pruning).")
doc.add_page_break()

# 5. REFLECTION
add_heading(doc, "5. רפלקציה / סיכום אישי", level=1)
add_rtl_paragraph(doc, "[כאן על התלמיד לכתוב רפלקציה אישית (לפחות חצי עמוד), הכוללת: כיצד הייתה העבודה על הפרויקט, מה קיבלת ממנה, אילו כלים אתה לוקח להמשך, מה הקשיים שעמדו בפניך, מה המסקנות שלך, ומה היית עושה אחרת לו היית מתחיל היום.]")
doc.add_page_break()

# 6. BIBLIOGRAPHY
add_heading(doc, "6. ביבליוגרפיה", level=1)
for i, ref in enumerate([
    "Mitchell, M. (1996). An Introduction to Genetic Algorithms. MIT Press.",
    "Goodfellow, I., Bengio, Y., & Courville, A. (2016). Deep Learning. MIT Press. https://www.deeplearningbook.org",
    "He, K., Zhang, X., Ren, S., & Sun, J. (2015). Delving Deep into Rectifiers. arXiv:1502.01852.",
    "Microsoft Docs. (2024). System.Threading.Interlocked. https://docs.microsoft.com/en-us/dotnet/api/system.threading.interlocked",
    "Wikipedia. (2024). Draughts. https://en.wikipedia.org/wiki/Draughts",
    "Silver, D., et al. (2017). Mastering Chess and Shogi by Self-Play. arXiv:1712.01815.",
], 1):
    add_rtl_paragraph(doc, f"{i}. {ref}")
doc.add_page_break()

# 7. APPENDICES
add_heading(doc, "7. נספחים", level=1)
add_heading(doc, "נספח א' - Leaky ReLU", level=2)
add_rtl_paragraph(doc, "פונקציית ה-Leaky ReLU מאפשרת שיפוע קטן לערכים שליליים - מונעת 'מוות' של נוירונים:")
add_code(doc, "private double LeakyReLU(double x) => x > 0 ? x : 0.01 * x;")
add_heading(doc, "נספח ב' - He Initialization", level=2)
add_code(doc, "double stdDev = Math.Sqrt(2.0 / rows);\nmatrix[i][j] = NextGaussian() * stdDev;")
add_heading(doc, "נספח ג' - Board State Encoding", level=2)
add_code(doc, "// '.' = empty, 'r'=Red, 'R'=Red King, 'b'=Black, 'B'=Black King\nchars[idx++] = piece == null ? '.' :\n               piece.Color == PieceColor.Red ?\n               (piece.Type == PieceType.King ? 'R' : 'r') :\n               (piece.Type == PieceType.King ? 'B' : 'b');")
add_heading(doc, "נספח ד' - Fitness Function (קטע קוד)", level=2)
add_code(doc, "double winRateWeight = 400 + (600 * experienceMultiplier); // 400->1000\nfitness += winRate * winRateWeight;\nfitness -= lossRate * (winRateWeight * 0.5);\nfitness += materialBalance * 5.0;\nfitness += Stats.KingsMade * 10.0;\nfitness += Stats.KingsCaptured * 20.0;\nfitness -= Stats.KingsLost * 25.0;")
add_heading(doc, "נספח ה' - Adaptive Mutation Rate", level=2)
add_code(doc, "double weakness = 1.0 - (avgParentFitness / bestFitness);\nrate += weakness * 0.15;\nif (generationsWithoutImprovement > 5)\n    rate += 0.05 * (generationsWithoutImprovement / 5.0);\nreturn Math.Min(rate, 0.5); // cap at 50%")

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
